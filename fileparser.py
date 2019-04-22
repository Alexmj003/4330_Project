import collections
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import docx
from wordcloud import WordCloud, STOPWORDS
import matplotlib.pyplot as plt

class FileParser:

    filename = input("Enter Filename:")
    def __init__(self,):
        input_file = None
        content = None# initial text of file
        content_split = None
        nltk.download('stopwords')
        self.stop_words = set(stopwords.words('english'))#stop word filtering
        word_limiter = 10
        result_file = None
        filter_words = []
        keywords = {}
        wordcounter = None
        self.topkeywords = []
        self.all_words = {}

    def readtxt(self,txt = str(filename),txt2 = str("Parsedwords.txt")):
        self.input_file = open(txt,'r')
        self.content = self.input_file.read()
        self.content_split = self.content.split(' ')#splitstring
        self.stop_words = set(stopwords.words('english'))#stop word filtering
        self.filter_words = []
        self.keywords = {}
        self.word_limiter = 10
        self.topkeywords = []
        for r in self.content_split:
            if not r in self.stop_words:
                self.filter_words.append(r)
        for word in self.filter_words:
            if word not in self.keywords:
                self.keywords[word] = 1
            else:
                self.keywords[word] += 1
        self.wordcounter = collections.Counter(self.keywords)
        for word, count in self.wordcounter.most_common(self.word_limiter):
            self.topkeywords.append(word)
        self.input_file.close()
        file = open(txt2,'a')
        file.truncate(0)
        for each in self.topkeywords:
            file.write(each + " ")
        file.close()

    def readDoc(self,txt = str(filename)):
        #I didn't change much of this because it worked
        self.input_file = docx.Document(txt)
        self.content = []
        #Obtained the list of stopwords including ''
        self.stop_words = set(stopwords.words('english'))
        self.stop_words.add('')
        #Condense the paragraphs and join them
        for para in self.input_file.paragraphs:
            self.content.append(para.text)
        text = "".join(self.content)
        self.content_split = text.split(' ')
        #If the word in not in stop_words add it to the dictionary
        for word in self.content_split:
            word = word.lower()
            if word not in self.stop_words:
                if word in self.all_words:
                    self.all_words[word] += 1
                else:
                    self.all_words[word] = 1
        #Create a list of the keys and values in the word dictionary to be sorted
        word_list = [[key, value] for key, value in self.all_words.items()]
        #sort the list in descending order
        word_list.sort(key=lambda x: x[1], reverse=True)
        #Save the top ten words
        self.topkeywords = [word[0] for word in word_list[:10]]


    def readdocx(self,txt = str(filename),txt2 = str("examplefile.txt")):
        self.input_file = docx.Document(txt)
        self.content = []
        self.stop_words = set(stopwords.words('english'))  # stop word filtering
        self.filter_words = []
        self.keywords = {}
        self.word_limiter = 10
        self.topkeywords = []
        for para in self.input_file.paragraphs:
            self.content.append(para.text)
        text = "".join(self.content)
        self.content_split = text.split(' ')
        for r in self.content_split:
            if not r in self.stop_words:
                self.filter_words.append(r)
        for word in self.filter_words:
            if word not in self.keywords:
                self.keywords[word] = 1
            else:
                self.keywords[word] += 1
        self.wordcounter = collections.Counter(self.keywords)
        for word, count in self.wordcounter.most_common(self.word_limiter):
            self.topkeywords.append(word)
        file = open(txt2,'a')
        file.truncate(0)
        for each in self.topkeywords:
            file.write(each + " ")
        file.close()

        

    def gettopkeywords(self):
        return self.topkeywords

    def wordcloud(self):
        string = ""
        for each in self.topkeywords:
            string += each + " "
        wordcloud = WordCloud().generate(string)
        return wordcloud

    def showwordcloud(self):
        string = ""
        for each in self.topkeywords:
            string += each + " "

        wordcloud = WordCloud().generate(string)
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show()

    def outputtxt(self,txt = str("")):
        file = open(txt,'a')
        file.truncate(0)
        self.filter_words = self.filter_words
        for each in self.filter_words:
            file.write(each + " ")
        file.close()

def __main__():
    document = FileParser()
    a = input('enter first file name : ')
    b = input('enter second file namee : ')
    document.readdocx(a,b)
    words = document.gettopkeywords()
    print(words)
    print(document.all_words)


if __name__ == '__main__':
    __main__()
