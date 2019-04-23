
import collections
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import docx
from wordcloud import WordCloud, STOPWORDS
import matplotlib.pyplot as plt

class FileParser:
    def __init__(self,):
        self.input_file = None
        self.content = None# initial text of file
        self.content_split = None
        self.stop_words = set(stopwords.words('english'))#stop word filtering
        self.word_limiter = 10
        self.result_file = None
        self.filter_words = []
        self.keywords = {}
        self.wordfreq = {}
        self.wordcounter = None
        self.topkeywords = []

    def readtxt(self,txt = str('PracticeFile.txt'), txt2 = str('FinalFile.txt')):
        self.input_file = open(txt,'r')
        self.content = self.input_file.read()
        self.content_split = self.content.split(' ')#splitstring
        for r in self.content_split:
            if not r in self.stop_words:
                self.filter_words.append(r)
        for word in self.filter_words:
            if word not in self.keywords:
                self.keywords[word] = 1
            else:
                self.keywords[word] += 1
        self.wordcounter = collections.Counter(self.keywords)
        for word, count in self.wordcounter.most_common(self.word_limiter):
            self.topkeywords.append(word)
            self.wordfreq[word] = count
        self.input_file.close()
        file = open(txt2,'a')
        file.truncate(0)
        for each in self.topkeywords:
            file.write(each + " " )
        file.close()

    def readdocx(self,txt = str('FileParser.docx'),txt2 = str('FinalFile.txt')):
        self.input_file = docx.Document(txt)
        self.stop_words = set(stopwords.words('english'))  # stop word filtering
        for para in self.input_file.paragraphs:
            self.content.append(para.text)
        text = "".join(self.content)
        self.content_split = text.split(' ')
        for r in self.content_split:
            if not r in self.stop_words:
                self.filter_words.append(r)
        for word in self.filter_words:
            if word not in self.keywords:
                self.keywords[word] = 1
            else:
                self.keywords[word] += 1
        self.wordcounter = collections.Counter(self.keywords)
        for word, count in self.wordcounter.most_common(self.word_limiter):
            self.topkeywords.append(word)
            self.wordfreq[word] = count
        file = open(txt2, 'a')
        file.truncate(0)
        for each in self.topkeywords:
            file.write(each + " ")
        file.close()

    def get_frequency(self):
        return self.wordfreq

    def gettopkeywords(self):
        return self.topkeywords

    def wordcloud(self):
        string = ""
        for each in self.topkeywords:
            string += each + " "

        wordcloud = WordCloud().generate(string)
        return wordcloud

    def showwordcloud(self):
        string = ""
        for each in self.topkeywords:
            string += each + " "

        wordcloud = WordCloud().generate(string)
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show()

def __main__():
    Document = FileParser()
    Document.readtxt()
    words = Document.gettopkeywords()
    print("HELLO")
    print(type(words))
    for word in words:
        print(word)

if __name__ == '__main__':
    __main__()
'''
doc = docx.Document("/Users/idky/PycharmProjects/BasicProject/untitled/Corn, Cows, and Cash_word FINAL DRAFT.docx")
fo = open(input("File: "), "r") #open file
content = fo.read() #readfile
contentSplit = content.split(' ') #splitString
stop_words = set(stopwords.words('english'))#stop word filtering
appendFile = open('filteredtext.txt', 'a')  # new file of filtered words
appendFile.truncate(0)#Clears Temp file of txt
appendFile.close() #closes The File
for r in contentSplit:#loop
    if not r in stop_words:#filter stop words
        appendFile = open('filteredtext.txt', 'a') # opens the file again
        appendFile.write(" "+r) # append to file
        appendFile.close() # close file

FilterFile = open("filteredtext.txt","r") #split
FilteredContent = FilterFile.read() #filteredcontent
FilteredContentSplit = FilteredContent.split(' ') #split content

wordcount = {} # list of key words
for word in FilteredContentSplit: #loop for gathering key words
    if word not in wordcount: #nested loop
        wordcount[word] = 1 #algoritm
    else:
        wordcount[word] += 1 #algorithm


WordLimiter = 10 #limit for keywords
word_counter = collections.Counter(wordcount) #collections class for counting words
FinalFile = open("FinalFile.txt","w") #Opens file to be passed to the wordcloud

for word, count in word_counter.most_common(WordLimiter): # writes words to file
    FinalFile.write(word + " ")

FilterFile.close()#cleanup
fo.close()#cleanup
FinalFile.close()#cleanup
'''
